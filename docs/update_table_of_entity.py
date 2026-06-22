from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


SOURCE = Path(__file__).with_name("Table of entity.original.docx")
OUTPUT = Path(__file__).with_name("Table of entity.updated.docx")


ROWS_TO_ADD = {
    0: [
        ["created_at", "TIMESTAMP", "", "No", "", "ວັນແລະເວລາສ້າງຂໍ້ມູນ"],
        ["is_active", "TINYINT", "1", "No", "", "ສະຖານະເປີດໃຊ້ຂໍ້ມູນຄົນເຈັບ"],
        ["deleted_at", "DATETIME", "", "Yes", "", "ວັນແລະເວລາລຶບຂໍ້ມູນແບບ Soft Delete"],
    ],
    1: [
        ["document_no", "VARCHAR", "40", "Yes", "UQ", "ເລກທີເອກະສານໃບສັ່ງກວດ"],
        ["billing_no", "VARCHAR", "40", "Yes", "UQ", "ເລກທີໃບແຈ້ງໜີ້"],
    ],
    2: [
        ["is_active", "TINYINT", "1", "No", "", "ສະຖານະເປີດໃຊ້ປະເພດການກວດ"],
        ["deleted_at", "DATETIME", "", "Yes", "", "ວັນແລະເວລາລຶບຂໍ້ມູນແບບ Soft Delete"],
    ],
    3: [
        ["created_at", "TIMESTAMP", "", "No", "", "ວັນແລະເວລາສ້າງຂໍ້ມູນ"],
        ["is_active", "TINYINT", "1", "No", "", "ສະຖານະເປີດໃຊ້ບັນຊີພະນັກງານ"],
        ["deleted_at", "DATETIME", "", "Yes", "", "ວັນແລະເວລາລຶບຂໍ້ມູນແບບ Soft Delete"],
    ],
    4: [
        ["called_at", "DATETIME", "", "Yes", "", "ວັນແລະເວລາທີ່ເອີ້ນຄິວ"],
    ],
    5: [
        ["result_image_url", "VARCHAR", "255", "Yes", "", "ທີ່ຢູ່ໄຟລ໌ຮູບພາບຜົນການກວດ"],
        ["report_no", "VARCHAR", "40", "Yes", "UQ", "ເລກທີລາຍງານຜົນການກວດ"],
    ],
    6: [
        ["status", "VARCHAR", "20", "Yes", "", "ສະຖານະການຊຳລະເງິນ"],
        ["adjustment_reason", "TEXT", "", "Yes", "", "ເຫດຜົນໃນການປັບປຸງລາຍການຊຳລະ"],
        ["adjusted_by", "INT", "", "Yes", "", "ລະຫັດພະນັກງານຜູ້ປັບປຸງລາຍການ"],
        ["adjusted_at", "DATETIME", "", "Yes", "", "ວັນແລະເວລາປັບປຸງລາຍການ"],
        ["receipt_no", "VARCHAR", "40", "Yes", "UQ", "ເລກທີໃບຮັບເງິນ"],
    ],
}


def replace_cell_text_preserving_format(cell, value):
    text_nodes = cell._tc.xpath(".//w:t")
    if not text_nodes:
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.text = value
        return

    text_nodes[0].text = value
    if value.startswith(" ") or value.endswith(" "):
        text_nodes[0].set(qn("xml:space"), "preserve")
    else:
        text_nodes[0].attrib.pop(qn("xml:space"), None)

    for text_node in text_nodes[1:]:
        text_node.text = ""


def append_cloned_row(table, values):
    template_row = table.rows[-1]
    cloned_tr = deepcopy(template_row._tr)
    table._tbl.append(cloned_tr)
    new_row = table.rows[-1]

    if len(new_row.cells) != len(values):
        raise ValueError(
            f"Expected {len(new_row.cells)} values but received {len(values)}"
        )

    for cell, value in zip(new_row.cells, values):
        replace_cell_text_preserving_format(cell, value)


def main():
    document = Document(SOURCE)

    if len(document.tables) != 7:
        raise ValueError(f"Expected 7 tables, found {len(document.tables)}")

    for table_index, new_rows in ROWS_TO_ADD.items():
        table = document.tables[table_index]
        existing_fields = {
            row.cells[0].text.strip()
            for row in table.rows[2:]
            if row.cells
        }
        for values in new_rows:
            if values[0] not in existing_fields:
                append_cloned_row(table, values)
                existing_fields.add(values[0])

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
