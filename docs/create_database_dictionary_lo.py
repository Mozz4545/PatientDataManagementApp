from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("Radiology_Database_Dictionary_Lao.docx")

FONT = "Phetsarath OT"
NAVY = "123879"
BLUE = "2E74B5"
INK = "17213A"
MUTED = "5F6775"
HEADER_FILL = "E8EEF5"
ALT_FILL = "F7F9FC"
NEW_FILL = "FFF2CC"
WHITE = "FFFFFF"
GREEN = "E2F0D9"
RED = "FCE4E4"

TABLE_WIDTH = 9360
TABLE_INDENT = 120
TABLE_COLS = [1900, 1900, 1700, 3860]


ENTITIES = [
    {
        "name": "staff",
        "lao": "ພະນັກງານ",
        "rows": 2,
        "description": (
            "ເກັບຂໍ້ມູນຜູ້ໃຊ້ລະບົບ, ສິດການໃຊ້ງານ, "
            "ຕຳແໜ່ງ ແລະ ສະຖານະຂອງພະນັກງານ."
        ),
        "fields": [
            ("staff_id", "INT(11)", "PK, AUTO_INCREMENT", "ລະຫັດພະນັກງານ", ""),
            ("username", "VARCHAR(100)", "UQ, NOT NULL", "ຊື່ສຳລັບເຂົ້າລະບົບ", ""),
            ("password", "VARCHAR(255)", "NOT NULL", "ລະຫັດຜ່ານທີ່ Hash ແລ້ວ", "ບໍ່ຄວນເກັບລະຫັດຜ່ານແບບຂໍ້ຄວາມທຳມະດາ"),
            ("staff_name", "VARCHAR(255)", "NOT NULL", "ຊື່ແລະນາມສະກຸນພະນັກງານ", ""),
            ("role", "ENUM", "DEFAULT STAFF", "ສິດຜູ້ໃຊ້: ADMIN ຫຼື STAFF", ""),
            ("position", "VARCHAR(100)", "NULL", "ຕຳແໜ່ງວຽກ", ""),
            ("department", "VARCHAR(100)", "NULL", "ພະແນກທີ່ສັງກັດ", ""),
            ("phone", "VARCHAR(20)", "NULL", "ເບີໂທຕິດຕໍ່", ""),
            ("created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "ວັນແລະເວລາສ້າງຂໍ້ມູນ", ""),
            ("is_active", "TINYINT(1)", "NOT NULL, DEFAULT 1", "ສະຖານະເປີດໃຊ້ບັນຊີ", ""),
            ("deleted_at", "DATETIME", "NULL", "ເວລາທີ່ລຶບແບບ Soft Delete", ""),
        ],
    },
    {
        "name": "patients",
        "lao": "ຄົນເຈັບ",
        "rows": 2,
        "description": (
            "ເກັບຂໍ້ມູນສ່ວນຕົວ ແລະ ຂໍ້ມູນຕິດຕໍ່ຂອງຄົນເຈັບ. "
            "ຄົນເຈັບໜຶ່ງຄົນສາມາດມີໃບສັ່ງກວດໄດ້ຫຼາຍໃບ."
        ),
        "fields": [
            ("patient_id", "INT(11)", "PK, AUTO_INCREMENT", "ລະຫັດຄົນເຈັບ", ""),
            ("first_name", "VARCHAR(255)", "NOT NULL", "ຊື່ຄົນເຈັບ", ""),
            ("last_name", "VARCHAR(255)", "NOT NULL", "ນາມສະກຸນຄົນເຈັບ", ""),
            ("age", "INT(11)", "NULL", "ອາຍຸ", ""),
            ("gender", "ENUM", "NULL", "ເພດ: M, F ຫຼື Other", ""),
            ("phone", "VARCHAR(20)", "NULL", "ເບີໂທຄົນເຈັບ", ""),
            ("date_of_birth", "DATE", "NULL", "ວັນເດືອນປີເກີດ", ""),
            ("address", "TEXT", "NULL", "ທີ່ຢູ່", ""),
            ("emergency_phone", "VARCHAR(20)", "NULL", "ເບີໂທຕິດຕໍ່ສຸກເສີນ", ""),
            ("created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "ວັນແລະເວລາສ້າງຂໍ້ມູນ", ""),
            ("is_active", "TINYINT(1)", "NOT NULL, DEFAULT 1", "ສະຖານະເປີດໃຊ້ຂໍ້ມູນຄົນເຈັບ", "ໃໝ່ໃນ Schema: ຍັງບໍ່ໄດ້ Commit"),
            ("deleted_at", "DATETIME", "NULL", "ເວລາທີ່ລຶບແບບ Soft Delete", "ໃໝ່ໃນ Schema: ຍັງບໍ່ໄດ້ Commit"),
        ],
    },
    {
        "name": "exam_types",
        "lao": "ປະເພດການກວດ",
        "rows": 4,
        "description": (
            "ເກັບລາຍການກວດລັງສີ, ລາຍລະອຽດ ແລະ ລາຄາຂອງແຕ່ລະປະເພດ."
        ),
        "fields": [
            ("exam_type_id", "INT(11)", "PK, AUTO_INCREMENT", "ລະຫັດປະເພດການກວດ", ""),
            ("exam_name", "VARCHAR(255)", "NOT NULL", "ຊື່ປະເພດການກວດ", ""),
            ("description", "TEXT", "NULL", "ລາຍລະອຽດການກວດ", ""),
            ("price", "DECIMAL(10,2)", "DEFAULT 0.00", "ລາຄາການກວດ", ""),
            ("is_active", "TINYINT(1)", "NOT NULL, DEFAULT 1", "ສະຖານະເປີດໃຊ້ປະເພດການກວດ", ""),
            ("deleted_at", "DATETIME", "NULL", "ເວລາທີ່ລຶບແບບ Soft Delete", ""),
        ],
    },
    {
        "name": "order",
        "lao": "ໃບສັ່ງກວດ",
        "rows": 2,
        "description": (
            "ເປັນ Entity ຫຼັກທີ່ເຊື່ອມຄົນເຈັບ, ປະເພດການກວດ, "
            "ພະນັກງານ, ຄິວ, ການຊຳລະເງິນ ແລະ ຜົນການກວດ."
        ),
        "fields": [
            ("order_id", "INT(11)", "PK, AUTO_INCREMENT", "ລະຫັດໃບສັ່ງກວດ", ""),
            ("patient_id", "INT(11)", "FK, NOT NULL", "ອ້າງອີງໄປຫາ patients.patient_id", ""),
            ("exam_type_id", "INT(11)", "FK, NOT NULL", "ອ້າງອີງໄປຫາ exam_types.exam_type_id", ""),
            ("staff_id", "INT(11)", "FK, NOT NULL", "ພະນັກງານຜູ້ສ້າງໃບສັ່ງ", ""),
            ("order_date", "DATETIME", "NOT NULL", "ວັນແລະເວລາສັ່ງກວດ", ""),
            ("note", "TEXT", "NULL", "ໝາຍເຫດເພີ່ມເຕີມ", ""),
            ("status", "VARCHAR(50)", "DEFAULT PENDING", "ສະຖານະໃບສັ່ງກວດ", ""),
            ("document_no", "VARCHAR(40)", "UQ, NULL", "ເລກທີເອກະສານໃບສັ່ງ", ""),
            ("billing_no", "VARCHAR(40)", "UQ, NULL", "ເລກທີໃບແຈ້ງໜີ້", ""),
        ],
    },
    {
        "name": "queue",
        "lao": "ຄິວກວດ",
        "rows": 2,
        "description": (
            "ເກັບລຳດັບຄິວການກວດຕາມວັນທີ. "
            "ໃບສັ່ງກວດໜຶ່ງໃບມີໄດ້ໜຶ່ງຄິວ."
        ),
        "fields": [
            ("queue_id", "INT(11)", "PK, AUTO_INCREMENT", "ລະຫັດຄິວ", ""),
            ("order_id", "INT(11)", "FK, UQ, NOT NULL", "ອ້າງອີງໄປຫາ order.order_id", "ກຳນົດ UQ ເພື່ອໃຫ້ໜຶ່ງໃບສັ່ງມີໜຶ່ງຄິວ"),
            ("queue_no", "INT(11)", "NOT NULL", "ເລກລຳດັບຄິວ", "UQ ຮ່ວມກັບ queue_date"),
            ("queue_date", "DATE", "NOT NULL", "ວັນທີຂອງຄິວ", "UQ ຮ່ວມກັບ queue_no"),
            ("status", "VARCHAR(50)", "DEFAULT WAITING", "ສະຖານະຄິວ", ""),
            ("called_at", "DATETIME", "NULL", "ເວລາທີ່ເອີ້ນຄິວ", ""),
        ],
    },
    {
        "name": "payment",
        "lao": "ການຊຳລະເງິນ",
        "rows": 1,
        "description": (
            "ເກັບຂໍ້ມູນການຮັບເງິນ, ໃບຮັບເງິນ "
            "ແລະ ປະຫວັດການປັບປຸງການຊຳລະ."
        ),
        "fields": [
            ("payment_id", "INT(11)", "PK, AUTO_INCREMENT", "ລະຫັດການຊຳລະເງິນ", ""),
            ("order_id", "INT(11)", "FK, UQ, NOT NULL", "ອ້າງອີງໄປຫາ order.order_id", "ໜຶ່ງໃບສັ່ງມີການຊຳລະໄດ້ບໍ່ເກີນໜຶ່ງລາຍການ"),
            ("staff_id", "INT(11)", "FK, NOT NULL", "ພະນັກງານຜູ້ຮັບຊຳລະ", ""),
            ("amount", "DECIMAL(10,2)", "NOT NULL", "ຈຳນວນເງິນ", ""),
            ("payment_date", "DATETIME", "NOT NULL", "ວັນແລະເວລາຊຳລະ", ""),
            ("payment_type", "VARCHAR(100)", "NOT NULL", "ຮູບແບບການຊຳລະ", ""),
            ("status", "VARCHAR(20)", "DEFAULT PAID, INDEX", "ສະຖານະການຊຳລະ", ""),
            ("adjustment_reason", "TEXT", "NULL", "ເຫດຜົນໃນການປັບປຸງລາຍການ", ""),
            ("adjusted_by", "INT(11)", "NULL", "ລະຫັດພະນັກງານຜູ້ປັບປຸງ", "ໃນ DB ປັດຈຸບັນຍັງບໍ່ສະແດງເປັນ FK"),
            ("adjusted_at", "DATETIME", "NULL", "ວັນແລະເວລາປັບປຸງ", ""),
            ("receipt_no", "VARCHAR(40)", "UQ, NULL", "ເລກທີໃບຮັບເງິນ", ""),
        ],
    },
    {
        "name": "result",
        "lao": "ຜົນການກວດ",
        "rows": 2,
        "description": (
            "ເກັບລາຍລະອຽດຜົນການກວດ, ຮູບພາບຜົນກວດ, "
            "ພະນັກງານຜູ້ບັນທຶກ ແລະ ເລກທີລາຍງານ."
        ),
        "fields": [
            ("result_id", "INT(11)", "PK, AUTO_INCREMENT", "ລະຫັດຜົນການກວດ", ""),
            ("order_id", "INT(11)", "FK, NOT NULL", "ອ້າງອີງໄປຫາ order.order_id", "ບໍ່ມີ UQ: ໜຶ່ງໃບສັ່ງອາດມີຫຼາຍຜົນກວດ"),
            ("staff_id", "INT(11)", "FK, NOT NULL", "ພະນັກງານຜູ້ບັນທຶກຜົນ", ""),
            ("result_detail", "TEXT", "NOT NULL", "ລາຍລະອຽດຜົນການກວດ", ""),
            ("result_date", "DATETIME", "NOT NULL", "ວັນແລະເວລາອອກຜົນ", ""),
            ("result_image_url", "VARCHAR(255)", "NULL", "ທີ່ຢູ່ໄຟລ໌ຮູບພາບຜົນກວດ", ""),
            ("report_no", "VARCHAR(40)", "UQ, NULL", "ເລກທີລາຍງານຜົນກວດ", ""),
        ],
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def apply_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, size=9, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, size=9, bold=False, color=INK, after=0, align=None):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.15
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def add_text(doc, text, size=10, bold=False, color=INK, after=6, align=None, italic=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    return paragraph


def add_table(doc, headers, rows, widths, note_column=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    apply_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        cell.text = value
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, HEADER_FILL)
        style_paragraph(
            cell.paragraphs[0],
            size=8.5,
            bold=True,
            color=NAVY,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        row_has_new = (
            note_column is not None
            and "ໃໝ່" in str(values[note_column])
        )
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_has_new:
                set_cell_shading(cells[idx], NEW_FILL)
            elif row_index % 2:
                set_cell_shading(cells[idx], ALT_FILL)
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 1) else WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(cells[idx].paragraphs[0], size=8.2, color=INK, align=align)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("ໜ້າ ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 8),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = "ລະບົບຈັດການຂໍ້ມູນຄົນເຈັບພະແນກລັງສີ | Database Dictionary"
    style_paragraph(header_p, size=8.2, color=MUTED)

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])

    add_text(doc, "ພະຈະນານຸກົມຖານຂໍ້ມູນ", size=24, bold=True, color=NAVY, after=4)
    add_text(
        doc,
        "ລະບົບຈັດການຂໍ້ມູນຄົນເຈັບພະແນກລັງສີ",
        size=14,
        bold=True,
        color=BLUE,
        after=8,
    )
    add_text(
        doc,
        "ກວດສອບຈາກຖານຂໍ້ມູນ radiology_db ໃນວັນທີ 21 ມິຖຸນາ 2026",
        size=9.5,
        color=MUTED,
        after=16,
    )

    add_heading(doc, "1. ຄຳອະທິບາຍສັນຍາລັກ", level=1)
    legend_rows = [
        ("PK", "Primary Key", "ຄີຫຼັກສຳລັບລະບຸແຕ່ລະແຖວບໍ່ໃຫ້ຊ້ຳ"),
        ("FK", "Foreign Key", "ຄີອ້າງອີງໄປຫາຕາຕະລາງອື່ນ"),
        ("UQ", "Unique Constraint", "ຂໍ້ກຳນົດທີ່ຫ້າມຄ່າຊ້ຳ"),
        ("NULL", "Nullable", "ອະນຸຍາດໃຫ້ບໍ່ມີຄ່າ"),
        ("AUTO_INCREMENT", "Automatic number", "ຖານຂໍ້ມູນສ້າງເລກລຳດັບໃຫ້ອັດຕະໂນມັດ"),
        ("Soft Delete", "Logical deletion", "ບັນທຶກເວລາລຶບໂດຍບໍ່ລຶບແຖວອອກຈິງ"),
    ]
    add_table(doc, ["ສັນຍາລັກ", "ຊື່ເຕັມ", "ຄວາມໝາຍ"], legend_rows, [1600, 2500, 5260])

    add_heading(doc, "2. ສະຫຼຸບ Entity ທັງໝົດ", level=1)
    summary_rows = [
        (entity["name"], entity["lao"], entity["rows"], len(entity["fields"]))
        for entity in ENTITIES
    ]
    add_table(
        doc,
        ["ຊື່ຕາຕະລາງ", "ຊື່ພາສາລາວ", "ຈຳນວນແຖວ", "ຈຳນວນ Field"],
        summary_rows,
        [2400, 3260, 1700, 2000],
    )
    add_text(
        doc,
        "ຖານຂໍ້ມູນຈິງມີ 7 ຕາຕະລາງ. ຈຳນວນແຖວແມ່ນຄ່າໃນເວລາກວດສອບ.",
        size=9,
        color=MUTED,
        after=4,
        italic=True,
    )

    for index, entity in enumerate(ENTITIES, start=1):
        doc.add_page_break()
        add_heading(doc, f"3.{index} Entity: {entity['name']} - {entity['lao']}", level=1)
        add_text(doc, entity["description"], size=10, after=8)
        add_text(
            doc,
            f"ຈຳນວນຂໍ້ມູນປັດຈຸບັນ: {entity['rows']} ແຖວ",
            size=9,
            bold=True,
            color=MUTED,
            after=6,
        )
        rows = [(field, data_type, key, meaning, note) for field, data_type, key, meaning, note in entity["fields"]]
        add_table(
            doc,
            ["Field", "ຊະນິດຂໍ້ມູນ", "Key / Constraint", "ຄວາມໝາຍ ແລະ ໝາຍເຫດ"],
            [(a, b, c, f"{d}\n{e}" if e else d) for a, b, c, d, e in rows],
            TABLE_COLS,
            note_column=3,
        )

    doc.add_page_break()
    add_heading(doc, "4. ຄວາມສຳພັນລະຫວ່າງ Entity", level=1)
    relation_rows = [
        ("patients", "1 : N", "order", "ຄົນເຈັບໜຶ່ງຄົນມີຫຼາຍໃບສັ່ງກວດ"),
        ("exam_types", "1 : N", "order", "ປະເພດການກວດໜຶ່ງປະເພດຖືກໃຊ້ໃນຫຼາຍໃບສັ່ງ"),
        ("staff", "1 : N", "order", "ພະນັກງານໜຶ່ງຄົນສ້າງຫຼາຍໃບສັ່ງ"),
        ("order", "1 : 1", "queue", "ໃບສັ່ງກວດໜຶ່ງໃບມີໜຶ່ງຄິວ"),
        ("order", "1 : 0..1", "payment", "ໃບສັ່ງກວດອາດຍັງບໍ່ມີ ຫຼື ມີການຊຳລະໜຶ່ງລາຍການ"),
        ("order", "1 : N", "result", "Schema ປັດຈຸບັນອະນຸຍາດໃຫ້ໜຶ່ງໃບສັ່ງມີຫຼາຍຜົນກວດ"),
        ("staff", "1 : N", "payment", "ພະນັກງານໜຶ່ງຄົນຮັບຊຳລະຫຼາຍລາຍການ"),
        ("staff", "1 : N", "result", "ພະນັກງານໜຶ່ງຄົນບັນທຶກຫຼາຍຜົນກວດ"),
    ]
    add_table(
        doc,
        ["Entity ຕົ້ນທາງ", "ຄວາມສຳພັນ", "Entity ປາຍທາງ", "ຄຳອະທິບາຍ"],
        relation_rows,
        [1900, 1500, 1900, 4060],
    )

    add_heading(doc, "5. ໝາຍເຫດລາຍການທີ່ເພີ່ມໃໝ່", level=1)
    change_rows = [
        (
            "Schema",
            "patients.is_active",
            "ເພີ່ມຄໍລຳສະຖານະເປີດໃຊ້ຂໍ້ມູນຄົນເຈັບ",
            "ມີຢູ່ໃນ DB ຈິງ ແຕ່ການປ່ຽນແປງ migrate.js ຍັງບໍ່ Commit",
        ),
        (
            "Schema",
            "patients.deleted_at",
            "ເພີ່ມຄໍລຳສຳລັບ Soft Delete",
            "ມີຢູ່ໃນ DB ຈິງ ແຕ່ການປ່ຽນແປງ migrate.js ຍັງບໍ່ Commit",
        ),
        (
            "Data",
            "patients ID 2",
            "ຂໍ້ມູນຄົນເຈັບເພີ່ມ 1 ແຖວ",
            "ເພີ່ມຫຼັງ Backup ວັນທີ 17 ມິຖຸນາ 2026",
        ),
        (
            "Data",
            "order ID 2",
            "ໃບສັ່ງກວດເພີ່ມ 1 ແຖວ",
            "ເຊື່ອມກັບ patients ID 2",
        ),
        (
            "Data",
            "queue ID 2",
            "ຄິວກວດເພີ່ມ 1 ແຖວ",
            "ເຊື່ອມກັບ order ID 2",
        ),
        (
            "Data",
            "result ID 2",
            "ຜົນການກວດເພີ່ມ 1 ແຖວ",
            "ເຊື່ອມກັບ order ID 2",
        ),
        (
            "Migration",
            "audit_logs",
            "ພົບຄຳສັ່ງ CREATE TABLE ໃນ migrate.js",
            "ຍັງບໍ່ພົບຕາຕະລາງນີ້ໃນ DB ຈິງ; ຍັງບໍ່ນັບເປັນ Entity ປັດຈຸບັນ",
        ),
    ]
    add_table(
        doc,
        ["ປະເພດ", "ລາຍການ", "ສິ່ງທີ່ເພີ່ມ", "ສະຖານະ / ໝາຍເຫດ"],
        change_rows,
        [1300, 2000, 2860, 3200],
        note_column=3,
    )

    add_text(
        doc,
        "ໝາຍເຫດ: ເອກະສານນີ້ສະແດງໂຄງສ້າງ ແລະ ຈຳນວນແຖວເທົ່ານັ້ນ; "
        "ບໍ່ໄດ້ເປີດເຜີຍຊື່ຄົນເຈັບ, ລະຫັດຜ່ານ ຫຼື ຂໍ້ມູນສ່ວນຕົວ.",
        size=9,
        color=MUTED,
        after=0,
        italic=True,
    )

    doc.core_properties.title = "ພະຈະນານຸກົມຖານຂໍ້ມູນ Radiology"
    doc.core_properties.subject = "Database entities and schema in Lao"
    doc.core_properties.author = "Radiology Patient Management System"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
