from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "UserManual_Lao"
SHOT_DIR = OUT_DIR / "screenshots"
OUT_FILE = OUT_DIR / "ບົດທີ 4_ຄູ່ມືການໃຊ້ງານລະບົບ_ສະບັບສົມບູນ_2026.docx"
TEMPLATE_FILE = Path(r"D:\Radiology Data Sysytem\ບົດທີ 4 (ຄຼ່ມືການໃຊ້ງານ).docx")
FONT = "Saysettha OT"
NAVY = RGBColor(18, 56, 121)
GRAY = RGBColor(92, 91, 110)
LIGHT = "F2F6FB"

def set_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

doc = Document(TEMPLATE_FILE)
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)
sec = doc.sections[0]
sec.page_width = Inches(8.27)
sec.page_height = Inches(11.69)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(12)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.0

for name in ["Title", "Subtitle"]:
    if name not in [style.name for style in styles]:
        styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

for name, size, color, before, after in [
    ("Title", 16, RGBColor(0, 0, 0), 0, 8),
    ("Subtitle", 14, RGBColor(0, 0, 0), 0, 8),
    ("Heading 1", 14, RGBColor(0, 0, 0), 12, 6),
    ("Heading 2", 14, RGBColor(0, 0, 0), 5, 2),
    ("Heading 3", 12, RGBColor(0, 0, 0), 5, 2),
]:
    s = styles[name]
    s.font.name = FONT
    s._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    s._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    s.font.size = Pt(size)
    s.font.color.rgb = color
    s.font.bold = name in {"Title", "Heading 1", "Heading 2", "Heading 3"}
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

for style_name in ["List Bullet", "List Number"]:
    if style_name not in [style.name for style in styles]:
        created_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        created_style.base_style = styles["List Paragraph"]
    s = styles[style_name]
    s.font.name = FONT
    s._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    s._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    s.font.size = Pt(12)
    s.paragraph_format.left_indent = Inches(0)
    s.paragraph_format.first_line_indent = Inches(0.59)
    s.paragraph_format.space_after = Pt(5)
    s.paragraph_format.line_spacing = 1.0

header = sec.header
header.paragraphs[0].text = ""
footer = sec.footer
footer.paragraphs[0].text = ""

def add_title(text, subtitle=None):
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    if subtitle:
        sp = doc.add_paragraph(style="Subtitle")
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.add_run(subtitle)

def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_numbering(p, 4)
        p.add_run(item)

def add_steps(items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_numbering(p, 11)
        p.add_run(item)

def set_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_element)

def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run("ໝາຍເຫດ: " + text)
    set_font(run, 12, RGBColor(0, 0, 0), bold=True)

def add_image(filename, caption):
    path = SHOT_DIR / filename
    if not path.exists():
        return
    with Image.open(path) as img:
        ratio = img.height / img.width
    width = 6.27
    height = width * ratio
    if height > 6.9:
        height = 6.9
        width = height / ratio
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    picture = p.add_run().add_picture(str(path), width=Inches(width), height=Inches(height))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(7)
    run = cp.add_run(caption)
    set_font(run, 12, RGBColor(0, 0, 0))

def page_break():
    doc.add_page_break()

def add_page_section(number, title, purpose, image, controls, functions=None, fields=None, notes=None):
    doc.add_heading(f"{number}. {title}", level=1)
    p = doc.add_paragraph()
    run = p.add_run("ຈຸດປະສົງ: ")
    set_font(run, 12, RGBColor(0, 0, 0), bold=True)
    p.add_run(purpose)
    add_image(image, f"ຮູບທີ {number}: {title}")
    doc.add_heading("ປຸ່ມ ແລະ ສ່ວນຄວບຄຸມ", level=2)
    add_bullets(controls)
    if fields:
        doc.add_heading("ຊ່ອງຂໍ້ມູນ", level=2)
        add_bullets(fields)
    if functions:
        doc.add_heading("ການເຮັດວຽກຂອງໜ້າ", level=2)
        add_steps(functions)
    if notes:
        add_note(notes)
    page_break()

# Chapter opening, following the supplied Chapter 4 template.
doc.add_paragraph().paragraph_format.space_after = Pt(54)
add_title(
    "ບົດທີ 4\nການນຳໃຊ້ງານຂອງລະບົບ",
    "ລະບົບຈັດການຂໍ້ມູນຄົນເຈັບ ພະແນກລັງສີ - ໂຮງໝໍ 103",
)
page_break()

doc.add_heading("ສາລະບານ", level=1)
toc = [
    "1. ໜ້າເຂົ້າລະບົບ",
    "2. ໂຄງສ້າງເມນູ ແລະ ສິດຜູ້ໃຊ້",
    "3. ໜ້າຫຼັກ",
    "4. ຂໍ້ມູນຄົນເຈັບ",
    "5. ປະຫວັດ ແລະ ແກ້ໄຂຄົນເຈັບ",
    "6. ຈັດການໃບສັ່ງກວດ",
    "7. ສ້າງໃບສັ່ງກວດໃໝ່",
    "8. ຈັດການຄິວ ແລະ ຈໍສະແດງ",
    "9. ການຊຳລະເງິນ",
    "10. ຜົນກວດ",
    "11. ປະເພດການກວດ",
    "12. ລາຍງານ",
    "13. ຂໍ້ມູນພະນັກງານ",
    "14. ປະຫວັດການໃຊ້ງານ (Audit Log)",
    "15. ຄວາມປອດໄພ ແລະ ການສຳຮອງຂໍ້ມູນ",
    "16. ສະຖານະວຽກ ແລະ ຂໍ້ຄວນລະວັງ",
]
add_bullets(toc)
add_note("ປຸ່ມທີ່ປ່ຽນແປງ ຫຼື ລົບຂໍ້ມູນຈະມີໜ້າຢືນຢັນ. ຄູ່ມືນີ້ອະທິບາຍຜົນຂອງປຸ່ມໂດຍບໍ່ລົບຂໍ້ມູນຈິງ.")
page_break()

add_page_section(
    "1", "ໜ້າເຂົ້າລະບົບ",
    "ກວດສອບຊື່ຜູ້ໃຊ້ ແລະ ລະຫັດຜ່ານ ກ່ອນເຂົ້າໃຊ້ລະບົບ.",
    "01-login.png",
    [
        "ຊ່ອງ “ລະຫັດຜູ້ໃຊ້”: ປ້ອນ username ຂອງພະນັກງານ.",
        "ຊ່ອງ “ລະຫັດຜ່ານ”: ປ້ອນ password; ຕົວອັກສອນຈະຖືກປິດບັງ.",
        "ປຸ່ມ “ຢືນຢັນ”: ສົ່ງຂໍ້ມູນເຂົ້າລະບົບ. ຖ້າສຳເລັດຈະໄປໜ້າຫຼັກ; ຖ້າຜິດຈະສະແດງຂໍ້ຄວາມສີແດງ.",
        "ປຸ່ມ “ເປີດຈໍສະແດງຄິວ”: ເຂົ້າໜ້າຈໍຄິວໄດ້ໂດຍບໍ່ຕ້ອງເຂົ້າລະບົບ.",
    ],
    [
        "ລະບົບກວດ session cookie ທີ່ປອດໄພກ່ອນເຂົ້າທຸກໜ້າທີ່ຕ້ອງ login.",
        "ລະບົບກວດວ່າທັງສອງຊ່ອງບໍ່ວ່າງ.",
        "ຫຼັງຈາກຢືນຢັນສຳເລັດ ເຊີບເວີຈະສົ່ງ HttpOnly Cookie; JavaScript ໃນ browser ບໍ່ສາມາດອ່ານ token ໄດ້.",
        "ຖ້າບັນຊີຖືກປິດໃຊ້, session ໝົດອາຍຸ ຫຼື role ຖືກປ່ຽນ ລະບົບຈະກວດຄືນຈາກຖານຂໍ້ມູນ ແລະ ປະຕິເສດການເຂົ້າ.",
    ],
)

doc.add_heading("2. ໂຄງສ້າງເມນູ ແລະ ສິດຜູ້ໃຊ້", level=1)
doc.add_heading("ແຖບເທິງ", level=2)
add_bullets([
    "ສະແດງຊື່ລະບົບ, ຊື່ຜູ້ໃຊ້ ແລະ ບົດບາດ ADMIN/STAFF.",
    "ປຸ່ມ “ອອກຈາກລະບົບ”: ລຶບ session cookie, ລ້າງ cache ຂໍ້ມູນ ແລະ ກັບໄປໜ້າ login.",
    "ໃນຈໍຂະໜາດນ້ອຍ ປຸ່ມ “ເມນູ” ຈະເປີດ/ປິດລາຍການເມນູ.",
])
doc.add_heading("ເມນູຫຼັກ", level=2)
add_bullets([
    "ໜ້າຫຼັກ: ສະຫຼຸບວຽກປະຈຳວັນ.",
    "ຂໍ້ມູນຄົນເຈັບ: ຄົ້ນຫາ, ແກ້ໄຂ ແລະ ເບິ່ງປະຫວັດຄົນເຈັບ.",
    "ຄິວ: ເອີ້ນຄິວ ແລະ ຄວບຄຸມຈໍສະແດງ.",
    "ໃບສັ່ງກວດ: ສ້າງ, ເບິ່ງ ແລະ ຍົກເລີກໃບສັ່ງ.",
    "ການຊຳລະເງິນ: ອອກໃບແຈ້ງ, ຮັບເງິນ ແລະ ອອກໃບຮັບເງິນ.",
    "ຜົນກວດ: ບັນທຶກລາຍງານ ແລະ ຮູບຜົນກວດ.",
    "ປະເພດການກວດ: ຈັດການຊື່, ລາຍລະອຽດ ແລະ ລາຄາ.",
    "ຂໍ້ມູນພະນັກງານ: ເຫັນສະເພາະ ADMIN; ໃຊ້ຈັດການບັນຊີ ແລະ ສິດ.",
    "ລາຍງານ: ADMIN ເບິ່ງລາຍງານການຊຳລະ ແລະ ຜົນກວດ; STAFF ເບິ່ງໄດ້ສະເພາະລາຍງານຜົນກວດ.",
    "ປະຫວັດການໃຊ້ງານ: ເຫັນສະເພາະ ADMIN; ໃຊ້ກວດວ່າໃຜເຮັດຫຍັງ, ເວລາໃດ ແລະ ຈາກ IP ໃດ.",
])
doc.add_heading("ການນຳທາງໃນມືຖື", level=2)
add_image("18-mobile-menu.png", "ຮູບທີ 2.1: ເມນູທັງໝົດໃນມືຖື")
add_bullets([
    "ແຖບລຸ່ມສະແດງ 5 ເມນູທີ່ໃຊ້ປະຈຳ: ໜ້າຫຼັກ, ຄິວ, ໃບສັ່ງກວດ, ການຊຳລະເງິນ ແລະ ຜົນກວດ.",
    "ປຸ່ມ “ເມນູ” ດ້ານເທິງສະແດງເມນູທັງໝົດ. ສຳລັບ ADMIN ຈະມີພະນັກງານ, ລາຍງານ ແລະ ປະຫວັດການໃຊ້ງານ.",
])
doc.add_heading("ສິດຜູ້ໃຊ້", level=2)
add_bullets([
    "ADMIN: ໃຊ້ງານທຸກໜ້າ, ລົບ/ປິດໃຊ້ຂໍ້ມູນ, Void/Refund, ຈັດການພະນັກງານ, ລາຍງານການເງິນ ແລະ Audit Log.",
    "STAFF: ດຳເນີນວຽກປະຈຳວັນ ແລະ ເບິ່ງລາຍງານຜົນກວດ; ບໍ່ສາມາດເຂົ້າຂໍ້ມູນພະນັກງານ, ລາຍງານການເງິນ ຫຼື Audit Log ແມ່ນແຕ່ພິມ URL ໂດຍກົງ.",
])
page_break()

add_page_section(
    "3", "ໜ້າຫຼັກ",
    "ສະແດງພາບລວມຄິວ, ໃບສັ່ງ, ຜົນກວດ ແລະ ການຊຳລະ.",
    "02-dashboard.png",
    [
        "“ສ້າງໃບສັ່ງກວດ”: ເປີດຟອມສ້າງໃບສັ່ງໃໝ່.",
        "“ໂຫຼດໃໝ່”: ໂຫຼດ orders, queues, payments ແລະ results ໃໝ່ພ້ອມກັນ.",
        "ກ່ອງ “ຄິວລໍຖ້າ”: ນັບຄິວທີ່ຍັງລໍຖ້າ.",
        "ກ່ອງ “ໃບສັ່ງມື້ນີ້”: ນັບໃບສັ່ງທີ່ບໍ່ຖືກຍົກເລີກໃນມື້ປັດຈຸບັນ.",
        "ກ່ອງ “ລໍຖ້າບັນທຶກຜົນ”: ນັບໃບສັ່ງທີ່ຍັງບໍ່ມີຜົນກວດ.",
        "ກ່ອງ “ຄ້າງຊຳລະ”: ນັບລາຍການພ້ອມຈ່າຍແຕ່ຍັງບໍ່ມີການຊຳລະ.",
        "ຕາຕະລາງລຸ່ມ: ສະແດງໃບສັ່ງຫຼ້າສຸດສູງສຸດ 6 ລາຍການ.",
    ],
)

add_page_section(
    "4", "ຂໍ້ມູນຄົນເຈັບ",
    "ຄົ້ນຫາ, ເບິ່ງຜົນ, ເບິ່ງປະຫວັດ, ແກ້ໄຂ ແລະ ລົບຄົນເຈັບ.",
    "03-patients.png",
    [
        "“ສ້າງໃບສັ່ງກວດ”: ເປີດຟອມສ້າງໃບສັ່ງ.",
        "ຊ່ອງຄົ້ນຫາ: ຄົ້ນດ້ວຍ HN/ID, ຊື່ ຫຼື ເບີໂທ; ຕາຕະລາງປ່ຽນຕາມຄຳຄົ້ນ.",
        "“ໂຫຼດໃໝ່”: ດຶງລາຍຊື່ຄົນເຈັບໃໝ່.",
        "“ເບິ່ງຜົນກວດ (n)”: ເປີດກ່ອງປະຫວັດຜົນກວດ; ປຸ່ມຈະຖືກປິດຖ້າບໍ່ມີຜົນ.",
        "“ປະຫວັດລະອຽດ (n)”: ເປີດໜ້າປະຫວັດໃບສັ່ງທັງໝົດຂອງຄົນເຈັບ.",
        "“ແກ້ໄຂ”: ເປີດຟອມຂໍ້ມູນຄົນເຈັບ.",
        "“ລົບ” (ADMIN): ສະແດງຄຳຢືນຢັນ ແລະ ລົບ/ຊ່ອນຄົນເຈັບຈາກລາຍຊື່.",
        "“ປິດ” ໃນກ່ອງຜົນກວດ: ປິດ modal ແລະ ກັບຄືນຕາຕະລາງ.",
    ],
)

doc.add_heading("5. ປະຫວັດ ແລະ ແກ້ໄຂຄົນເຈັບ", level=1)
add_image("15-patient-history.png", "ຮູບທີ 5.1: ປະຫວັດໃບສັ່ງກວດ")
doc.add_heading("ໜ້າປະຫວັດ", level=2)
add_bullets([
    "“ແກ້ໄຂຂໍ້ມູນ”: ໄປຫາຟອມແກ້ໄຂຄົນເຈັບ.",
    "“ກັບຄືນ”: ກັບໄປລາຍຊື່ຄົນເຈັບ.",
    "ສ່ວນຂໍ້ມູນຄົນເຈັບ: ສະແດງ HN, ຊື່, ເພດ, ອາຍຸ, ວັນເກີດ, ເບີໂທ, ເບີສຸກເສີນ, ທີ່ຢູ່ ແລະ ວັນສ້າງ.",
    "ໃນແຕ່ລະໃບສັ່ງ ສະແດງສະຖານະ, ວັນທີ, ຜູ້ສ້າງ, ລາຄາ, ໝາຍເຫດ, ຜົນກວດ ແລະ ການຊຳລະ.",
    "“ເປີດຮູບຜົນກວດ”: ໂຫຼດຮູບຜົນກວດແບບປອດໄພແລ້ວສະແດງເຕັມຈໍ.",
    "“ປິດ” ໃນຮູບເຕັມຈໍ: ປິດຮູບ ແລະ ຄືນຄ່າ memory ຂອງຮູບ.",
])
add_image("16-patient-edit.png", "ຮູບທີ 5.2: ຟອມແກ້ໄຂຄົນເຈັບ")
doc.add_heading("ຟອມແກ້ໄຂ", level=2)
add_bullets([
    "ຂໍ້ມູນບັງຄັບ: ຊື່, ນາມສະກຸນ ແລະ ເພດ.",
    "ຂໍ້ມູນເພີ່ມເຕີມ: ອາຍຸ, ວັນເກີດ, ເບີໂທ, ເບີສຸກເສີນ ແລະ ທີ່ຢູ່.",
    "“ປະຫວັດ”: ກັບໄປໜ້າປະຫວັດຂອງຄົນນີ້.",
    "“ກັບຄືນ”: ກັບໄປລາຍຊື່ໂດຍບໍ່ບັນທຶກ.",
    "“ບັນທຶກການແກ້ໄຂ”: ກວດ validation, ອັບເດດຖານຂໍ້ມູນ, ໂຫຼດລາຍຊື່/ໃບສັ່ງ/ຜົນກວດໃໝ່ ແລະ ກັບໄປໜ້າຄົນເຈັບ.",
])
page_break()

add_page_section(
    "6", "ຈັດການໃບສັ່ງກວດ",
    "ເບິ່ງໃບສັ່ງກວດທັງໝົດ ແລະ ຍົກເລີກລາຍການທີ່ຍັງດຳເນີນງານ.",
    "04-orders.png",
    [
        "“ສ້າງໃບສັ່ງກວດ”: ໄປຫາຟອມສ້າງໃໝ່.",
        "“ໂຫຼດໃໝ່”: ດຶງລາຍການຈາກ API ໃໝ່.",
        "ປຸ່ມ “x” ທ້າຍແຖວ: ເປີດກ່ອງຢືນຢັນຍົກເລີກ; ຈະສະແດງສະເພາະໃບສັ່ງທີ່ຍັງລໍຖ້າ ຫຼື ລໍຖ້າບັນທຶກຜົນ.",
        "“ຍົກເລີກ” ໃນກ່ອງຢືນຢັນ: ປ່ຽນສະຖານະໃບສັ່ງເປັນຍົກເລີກ ແລະ ໂຫຼດ orders/queues ໃໝ່.",
        "“ກັບຄືນ” ໃນກ່ອງ: ປິດກ່ອງໂດຍບໍ່ປ່ຽນຂໍ້ມູນ.",
    ],
)

add_page_section(
    "7", "ສ້າງໃບສັ່ງກວດໃໝ່",
    "ເລືອກຄົນເຈັບເກົ່າ ຫຼື ສ້າງຄົນເຈັບໃໝ່ ແລ້ວອອກໃບສັ່ງກວດ.",
    "05-new-order.png",
    [
        "“ກັບຄືນ”: ກັບໄປໜ້າໃບສັ່ງກວດ.",
        "ຊ່ອງຄົ້ນຫາຄົນເຈັບເກົ່າ: ຄົ້ນດ້ວຍ HN/ID ຫຼື ຊື່; ສະແດງສູງສຸດ 8 ຜົນລັບ.",
        "ປຸ່ມລາຍຊື່ຜົນຄົ້ນຫາ: ເລືອກຄົນເຈັບ ແລະ ຕື່ມຟອມອັດຕະໂນມັດ.",
        "“ໃຊ້ຄົນເຈັບໃໝ່”: ຍົກເລີກການເລືອກຄົນເຈັບເກົ່າ ແລະ ລ້າງຂໍ້ມູນສ່ວນບຸກຄົນ.",
        "ຕົວເລືອກພະນັກງານ: ຄົ້ນດ້ວຍລະຫັດ, ຊື່, ຕຳແໜ່ງ ຫຼື ພະແນກ; ຄ່າເລີ່ມຕົ້ນແມ່ນຜູ້ທີ່ login.",
        "“ສ້າງໃບສັ່ງກວດ”: ກວດຟອມ, ສ້າງ patient ກ່ອນຖ້າເປັນຄົນໃໝ່, ຈາກນັ້ນສ້າງ order ແລະ queue ອັດຕະໂນມັດ.",
    ],
    fields=[
        "ຊື່, ນາມສະກຸນ, ອາຍຸ, ເພດ, ວັນເກີດ, ທີ່ຢູ່, ເບີໂທ ແລະ ເບີສຸກເສີນ.",
        "ປະເພດການກວດ: ລາຍຊື່ຈາກຖານຂໍ້ມູນ; ລາຄາຈະຜູກກັບປະເພດ.",
        "ພະນັກງານຜູ້ສ້າງໃບສັ່ງ ແລະ ໝາຍເຫດ.",
    ],
    notes="ຖ້າເລືອກຄົນເຈັບເກົ່າ ຊ່ອງຂໍ້ມູນບຸກຄົນຈະເປັນ read-only. ຖ້າເປັນຄົນໃໝ່ ຕ້ອງປ້ອນຊື່, ນາມສະກຸນ, ອາຍຸ, ທີ່ຢູ່ ແລະ ເບີໂທໃຫ້ຄົບ.",
)

doc.add_heading("8. ຈັດການຄິວ ແລະ ຈໍສະແດງ", level=1)
add_image("06-queues.png", "ຮູບທີ 8.1: ໜ້າຈັດການຄິວ")
add_bullets([
    "“ເອີ້ນຄິວ” ດ້ານເທິງ: ເລືອກຄິວລໍຖ້າລຳດັບຕໍ່ໄປຂອງວັນ ແລະ ປ່ຽນສະຖານະເປັນກຳລັງເອີ້ນ.",
    "ປຸ່ມ “ເອີ້ນ” ໃນແຖວ: ເອີ້ນຄິວທີ່ເລືອກໂດຍກົງ.",
    "ຫຼັງເອີ້ນ ຈະມີ modal ສະແດງເລກຄິວ, ຊື່ ແລະ ປະເພດກວດ.",
    "“ເປີດຈໍສະແດງ”: ເປີດຈໍຄິວໃນ window ໃໝ່ຂະໜາດ 1200 × 760.",
    "“ປິດ” ໃນ modal: ປິດກ່ອງການເອີ້ນ.",
    "“ໂຫຼດໃໝ່”: ໂຫຼດຄິວໃໝ່; ໜ້ານີ້ຍັງ refresh ອັດຕະໂນມັດທຸກ 5 ວິນາທີ.",
])
add_image("07-queue-display.png", "ຮູບທີ 8.2: ຈໍສະແດງຄິວສຳລັບຄົນເຈັບ")
add_bullets([
    "ສະແດງໂມງ ແລະ ວັນທີຕາມເຂດເວລາ Asia/Vientiane.",
    "ກ່ອງໃຫຍ່ສະແດງຄິວປັດຈຸບັນ, ຊື່, ປະເພດກວດ ແລະ ເວລາເອີ້ນ.",
    "ດ້ານຂວາສະແດງຄິວທີ່ເອີ້ນກ່ອນໜ້າ.",
    "ຈໍນີ້ refresh ອັດຕະໂນມັດທຸກ 2 ວິນາທີ ແລະ ບໍ່ມີປຸ່ມແກ້ໄຂຂໍ້ມູນ.",
])
page_break()

add_page_section(
    "9", "ການຊຳລະເງິນ",
    "ອອກໃບແຈ້ງຊຳລະ, ບັນທຶກການຈ່າຍ, ພິມໃບຮັບເງິນ ແລະ ປັບສະຖານະການຈ່າຍ.",
    "08-payments.png",
    [
        "ຊ່ອງຄົ້ນຫາ: ກອງດ້ວຍ ID, ຊື່, ເລກໃບຮັບ, ປະເພດກວດ ຫຼື ຊ່ອງທາງຈ່າຍ.",
        "“ໂຫຼດໃໝ່”: ໂຫຼດ orders ຫຼື payments ຕາມ tab ທີ່ເປີດ.",
        "tab “ລໍຖ້າຊຳລະ”: ສະແດງໃບສັ່ງທີ່ບໍ່ຖືກຍົກເລີກ ແລະ ຍັງບໍ່ຈ່າຍ.",
        "“ໃບແຈ້ງຊຳລະ”: ສ້າງເອກະສານສຳລັບພິມ; ບໍ່ແມ່ນໃບຮັບເງິນ.",
        "“ຈ່າຍເງິນ”: ເປີດຟອມຈ່າຍເມື່ອຜົນກວດສຳເລັດ; ຖ້າຍັງບໍ່ພ້ອມຈະສະແດງ “ລໍຖ້າຜົນກວດ”.",
        "ໃນຟອມຈ່າຍ: ເລືອກຜູ້ຮັບເງິນ, ຊ່ອງທາງ “ເງິນສົດ/ເງິນໂອນ”; ຖ້າເງິນສົດຈະຄຳນວນເງິນທອນ.",
        "“ຢືນຢັນການຈ່າຍ”: ກວດຈຳນວນເງິນ, ສ້າງ payment ແລະ ຍ້າຍໄປ tab ຈ່າຍແລ້ວ.",
        "tab “ຈ່າຍແລ້ວ”: ສະແດງລາຍການທີ່ສະຖານະ PAID.",
        "“ໃບຮັບເງິນ”: ເປີດຕົວຢ່າງໃບຮັບ; “ພິມໃບຮັບເງິນ” ເປີດ print dialog.",
        "“Void” (ADMIN): ຍົກເລີກທຸລະກຳທີ່ບັນທຶກຜິດ; ຕ້ອງລະບຸເຫດຜົນ.",
        "“Refund” (ADMIN): ບັນທຶກການຄືນເງິນ; ຕ້ອງລະບຸເຫດຜົນ.",
        "“ຍົກເລີກ”: ປິດຟອມຫຼື modal ໂດຍບໍ່ບັນທຶກ.",
    ],
    notes="ການ Void/Refund ບໍ່ລຶບປະຫວັດ; ລະບົບປ່ຽນສະຖານະ ແລະ ເກັບເຫດຜົນໄວ້.",
)

add_page_section(
    "10", "ຜົນກວດ",
    "ບັນທຶກ/ແກ້ໄຂຂໍ້ຄວາມຜົນກວດ, ແນບຮູບ ແລະ ພິມເອກະສານ.",
    "09-results.png",
    [
        "ຊ່ອງຄົ້ນຫາ: ກອງດ້ວຍ HN ຫຼື ຊື່ຄົນເຈັບ.",
        "“ໂຫຼດໃໝ່”: ໂຫຼດ orders ແລະ results ພ້ອມກັນ.",
        "“ກັບຄືນ”: ກັບໜ້າຫຼັກ.",
        "“ບັນທຶກຜົນ”: ເປີດຟອມສຳລັບ order ທີ່ຍັງບໍ່ມີ result.",
        "“ແກ້ໄຂ”: ເປີດຜົນກວດເກົ່າເພື່ອປ່ຽນຂໍ້ຄວາມ ຫຼື ຮູບ.",
        "“ເອກະສານ”: ສ້າງລາຍງານຜົນກວດສຳລັບພິມ; ຖ້າມີຮູບຈະແນບໃນເອກະສານ.",
        "“ເບິ່ງຮູບ”: ໂຫຼດຮູບທີ່ບັນທຶກແລ້ວແລະເປີດເຕັມຈໍ.",
        "“ເລືອກໄຟລ໌/ກົດເພື່ອປ່ຽນຮູບ”: ຮັບສະເພາະ JPG/PNG ຂະໜາດບໍ່ເກີນ 5 MB.",
        "“ຍົກເລີກຮູບໃໝ່”: ຖອນຮູບທີ່ເລືອກກ່ອນບັນທຶກ.",
        "“ລຶບຮູບ”: ໝາຍໃຫ້ລຶບຮູບເກົ່າເມື່ອກົດບັນທຶກ.",
        "ກົດຮູບ preview: ເປີດຮູບເຕັມຈໍ; “ປິດ” ຈະກັບຟອມ.",
        "“ບັນທຶກ”: ສົ່ງຂໍ້ຄວາມ ແລະ ຮູບແບບ multipart, ອັບເດດສະຖານະ workflow ແລະ ໂຫຼດລາຍການໃໝ່.",
    ],
)

add_page_section(
    "11", "ປະເພດການກວດ",
    "ຈັດການລາຍການບໍລິການກວດ ແລະ ລາຄາ; ການແກ້ໄຂຈຳກັດສະເພາະ ADMIN.",
    "10-exam-types.png",
    [
        "“ກັບຄືນ”: ກັບໜ້າຫຼັກ.",
        "“ບັນທຶກ”: ສ້າງປະເພດໃໝ່ ຫຼື ອັບເດດລາຍການທີ່ກຳລັງແກ້ໄຂ.",
        "“ແກ້ໄຂ”: ນຳຊື່, ລາຍລະອຽດ ແລະ ລາຄາຂອງແຖວມາໃສ່ຟອມ.",
        "“ຍົກເລີກການແກ້ໄຂ”: ລ້າງຟອມ ແລະ ກັບໂໝດເພີ່ມໃໝ່.",
        "“ລົບ”: ເປີດ modal ຢືນຢັນ; ປຸ່ມ “ລົບ” ໃນ modal ຈະຊ່ອນປະເພດອອກຈາກລາຍຊື່ໃໝ່ ແຕ່ປະຫວັດ order ເກົ່າຍັງຢູ່.",
        "“ຍົກເລີກ” ໃນ modal: ປິດ modal ໂດຍບໍ່ລົບ.",
    ],
    fields=[
        "ຊື່ປະເພດການກວດ: ບັງຄັບ.",
        "ລາຍລະອຽດ: ບໍ່ບັງຄັບ.",
        "ລາຄາ: ຕ້ອງເປັນເລກ 0 ຫຼື ຫຼາຍກວ່າ.",
    ],
)

add_page_section(
    "12", "ລາຍງານ",
    "ສະຫຼຸບ, ກອງ, ພິມ ແລະ ສົ່ງອອກລາຍງານ. ADMIN ເບິ່ງໄດ້ທັງການຊຳລະ ແລະ ຜົນກວດ; STAFF ເບິ່ງໄດ້ສະເພາະຜົນກວດ.",
    "13-reports-payments.png",
    [
        "ຊ່ອງຄົ້ນຫາ: ກອງຂໍ້ມູນຂອງ tab ທີ່ເປີດ.",
        "“ໂຫຼດໃໝ່”: ດຶງຂໍ້ມູນຂອງ tab ປັດຈຸບັນຈາກ API ໃໝ່.",
        "“ລ້າງຕົວກອງ”: ລ້າງຄຳຄົ້ນຫາ, ວັນທີ, ໄລຍະລາຍຮັບ ແລະ ຊ່ອງທາງຊຳລະ.",
        "ຕົວກອງວັນທີເລີ່ມ/ສິ້ນສຸດ: ໃຊ້ກັບລາຍງານຜົນກວດ.",
        "ຖ້າວັນທີເລີ່ມຫຼາຍກວ່າວັນທີສິ້ນສຸດ ລະບົບຈະເຕືອນ ແລະ ປິດປຸ່ມພິມ/ສົ່ງອອກ.",
        "ລາຍຮັບ “ທັງໝົດ/ລາຍວັນ/ລາຍເດືອນ”: ປ່ຽນຊ່ວງການສະຫຼຸບ.",
        "ຕົວກອງ “ທັງໝົດ/ເງິນສົດ/ເງິນໂອນ”: ກອງຕາມຊ່ອງທາງການຈ່າຍ.",
        "“ພິມເອກະສານ”: ສ້າງລາຍງານຕາມ tab ແລະ ຕົວກອງປັດຈຸບັນ ແລ້ວເປີດ print dialog.",
        "“ສົ່ງອອກ”: ສ້າງໄຟລ໌ CSV ຂອງ tab ປັດຈຸບັນ.",
        "tab “ລາຍງານການຊຳລະ”: ສະຫຼຸບລາຍຮັບຕາມຊ່ອງທາງ/ໄລຍະເວລາ ແລະ ລາຍການລະອຽດ.",
        "tab “ລາຍງານຜົນກວດ”: “ເບິ່ງຮູບ” ເປີດຮູບເຕັມຈໍ; “ເບິ່ງລາຍລະອຽດ” ເປີດ modal ຂໍ້ມູນ.",
        "“ປິດ” ໃນ modal ລາຍລະອຽດ/ຮູບ: ປິດໜ້າຕ່າງ.",
    ],
)
add_image("14-reports-results.png", "ຮູບທີ 12.2: ລາຍງານຜົນກວດ")

doc.add_heading("13. ຂໍ້ມູນພະນັກງານ", level=1)
add_image("11-staff.png", "ຮູບທີ 13.1: ລາຍຊື່ ແລະ ການຈັດການພະນັກງານ")
doc.add_heading("ໜ້າລາຍຊື່ພະນັກງານ", level=2)
add_bullets([
    "ຊ່ອງຄົ້ນຫາ: ຄົ້ນດ້ວຍລະຫັດ, ຊື່, username, ຕຳແໜ່ງ, ພະແນກ ຫຼື ເບີໂທ.",
    "“ໂຫຼດໃໝ່”: ດຶງລາຍຊື່ພະນັກງານໃໝ່.",
    "“ເພີ່ມພະນັກງານ”: ເປີດຟອມສ້າງບັນຊີ.",
    "“ພິມລາຍຊື່”: ສ້າງເອກະສານລາຍຊື່ຕາມຜົນຄົ້ນຫາປັດຈຸບັນ.",
    "“ສົ່ງອອກ CSV”: ດາວໂຫຼດລາຍຊື່ຕາມຜົນຄົ້ນຫາ.",
    "“ແກ້ໄຂ”: ເປີດຟອມແກ້ໄຂຂໍ້ມູນ ແລະ ສິດ.",
    "“ປິດການນຳໃຊ້”: ເປີດ modal ຢືນຢັນ; ບັນຊີຈະຖືກຊ່ອນແຕ່ປະຫວັດວຽກຍັງຢູ່.",
    "ບັນຊີທີ່ login ຢູ່ຈະສະແດງປ້າຍ “ບັນຊີທີ່ກຳລັງໃຊ້ງານ” ແລະ ບໍ່ສາມາດປິດໄດ້.",
])
add_image(
    "12-staff-new.png",
    "ຮູບທີ 13.2: ຟອມເພີ່ມ/ແກ້ໄຂພະນັກງານ",
)
doc.add_heading("ຟອມເພີ່ມ/ແກ້ໄຂ", level=2)
add_bullets(
    [
        "“ກັບຄືນ”: ກັບໄປໜ້າຂໍ້ມູນພະນັກງານ.",
        "“ບັນທຶກຂໍ້ມູນພະນັກງານ”: ກວດ validation ແລະ ສ້າງ/ອັບເດດບັນຊີ.",
        "ໃນໂໝດສ້າງໃໝ່ ລະຫັດຜ່ານບັງຄັບ ແລະ ຕ້ອງມີຢ່າງນ້ອຍ 6 ຕົວ.",
        "ໃນໂໝດແກ້ໄຂ ປ່ອຍລະຫັດຜ່ານວ່າງເພື່ອໃຊ້ລະຫັດເກົ່າ; ຖ້າປ້ອນຄ່າໃໝ່ ລະບົບຈະປ່ຽນລະຫັດແຍກຈາກຂໍ້ມູນທົ່ວໄປ.",
    ],
)
doc.add_heading("ຊ່ອງຂໍ້ມູນ", level=2)
add_bullets([
    "ຊື່-ນາມສະກຸນ ແລະ username: ບັງຄັບ.",
    "password: ບັງຄັບເມື່ອສ້າງໃໝ່.",
    "ສິດນຳໃຊ້: STAFF ຫຼື ADMIN.",
    "ຕຳແໜ່ງ, ພະແນກ ແລະ ເບີໂທ: ບໍ່ບັງຄັບ.",
])
page_break()

add_page_section(
    "14", "ປະຫວັດການໃຊ້ງານ (Audit Log)",
    "ເກັບ ແລະ ຄົ້ນຫາປະຫວັດການປ່ຽນແປງຂໍ້ມູນ; ເປີດໄດ້ສະເພາະ ADMIN.",
    "19-audit-logs.png",
    [
        "“ໂຫຼດໃໝ່”: ດຶງລາຍການຫຼ້າສຸດຈາກເຊີບເວີ.",
        "ຊ່ອງຄົ້ນຫາ: ຄົ້ນດ້ວຍຊື່ຜູ້ໃຊ້, ລາຍລະອຽດ, ລະຫັດຂໍ້ມູນ ຫຼື IP.",
        "ຕົວກອງການກະທຳ: ເຊັ່ນ ເຂົ້າລະບົບ, ເພີ່ມ, ແກ້ໄຂ, ຍົກເລີກ, ຊຳລະ, Void ແລະ Refund.",
        "ຕົວກອງປະເພດຂໍ້ມູນ: ຄົນເຈັບ, ພະນັກງານ, ໃບສັ່ງກວດ, ຄິວ, ຜົນກວດ, ການຊຳລະ ແລະ ປະເພດການກວດ.",
        "ວັນທີເລີ່ມ/ສິ້ນສຸດ: ຈຳກັດຊ່ວງເວລາທີ່ຕ້ອງການກວດສອບ.",
        "ປຸ່ມໜ້າກ່ອນ/ໜ້າຖັດໄປ: ປ່ຽນໜ້າເມື່ອມີຫຼາຍກວ່າ 20 ລາຍການ.",
    ],
    [
        "ລະບົບບັນທຶກສະເພາະ API ທີ່ປ່ຽນຂໍ້ມູນສຳເລັດ; ຄຳຂໍທີ່ validation ບໍ່ຜ່ານຈະບໍ່ສ້າງລາຍການຫຼອກ.",
        "ແຕ່ລະລາຍການມີວັນເວລາ, ຜູ້ໃຊ້, role, ການກະທຳ, ຂໍ້ມູນທີ່ກ່ຽວຂ້ອງ, ລາຍລະອຽດ ແລະ IP.",
        "STAFF ບໍ່ເຫັນເມນູ ແລະ ຖ້າເຂົ້າ /audit-logs ໂດຍກົງຈະຖືກສົ່ງກັບໜ້າຫຼັກ; API ຈະຕອບ 403.",
    ],
    notes="Audit Log ເປັນຫຼັກຖານກວດສອບ. ບໍ່ຄວນແກ້ໄຂ ຫຼື ລຶບຈາກຖານຂໍ້ມູນໂດຍກົງ.",
)

doc.add_heading("15. ຄວາມປອດໄພ ແລະ ການສຳຮອງຂໍ້ມູນ", level=1)
doc.add_heading("ຄວາມປອດໄພຂອງການເຂົ້າໃຊ້", level=2)
add_bullets([
    "JWT ຖືກເກັບໃນ HttpOnly, SameSite=Strict Cookie; ບໍ່ເກັບ token ໃນ localStorage.",
    "ທຸກ API ທີ່ປ້ອງກັນຈະກວດ token, ສະຖານະບັນຊີ ແລະ role ປັດຈຸບັນຈາກຖານຂໍ້ມູນ.",
    "Frontend guard ປ້ອງກັນ URL ໂດຍກົງ; Backend role guard ເປັນຊັ້ນປ້ອງກັນຫຼັກ.",
    "ໃນ production ຕ້ອງໃຊ້ HTTPS, JWT secret ທີ່ແຂງແຮງ, CORS ທີ່ລະບຸ origin ແລະ ປ່ຽນລະຫັດ ADMIN ເລີ່ມຕົ້ນ.",
])
doc.add_heading("ສ້າງ Backup", level=2)
add_steps([
    "ເປີດ Terminal ແລ້ວເຂົ້າໂຟນເດີ backend.",
    "ຮັນຄຳສັ່ງ npm run backup.",
    "ກວດໂຟນເດີ backend/backups/backup-YYYYMMDD-HHMMSS; ຕ້ອງມີ manifest.json, database.json ແລະ uploads.",
    "ເກັບສຳເນົາ Backup ໄວ້ນອກເຄື່ອງ Server ແລະ ຈຳກັດສິດຜູ້ເຂົ້າເຖິງ.",
])
doc.add_heading("Restore ແລະ ກວດສອບ", level=2)
add_steps([
    "ສຳຮອງຖານຂໍ້ມູນປັດຈຸບັນກ່ອນ Restore ທຸກຄັ້ງ.",
    "ຮັນ npm run migrate ເພື່ອກຽມ schema.",
    "ຮັນ npm run restore -- backup-YYYYMMDD-HHMMSS --yes.",
    "ຮັນ verify-restore.js ຫຼື ກວດຈຳນວນຂໍ້ມູນ, ຄວາມສຳພັນ ແລະ ຮູບຜົນກວດ.",
    "ທົດສອບ Login, ຄົ້ນຫາຄົນເຈັບ, ເບິ່ງຜົນກວດ ແລະ Audit Log ກ່ອນເປີດໃຊ້ງານ.",
])
add_note("Restore ຈະແທນທີ່ຂໍ້ມູນ ແລະ ຮູບຜົນກວດປັດຈຸບັນ. ຄວນເຮັດໂດຍ ADMIN/ຜູ້ດູແລ Server ເທົ່ານັ້ນ.")
page_break()

doc.add_heading("16. ສະຖານະວຽກ ແລະ ຂໍ້ຄວນລະວັງ", level=1)
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
table.columns[0].width = Inches(1.7)
table.columns[1].width = Inches(5.0)
header_properties = table.rows[0]._tr.get_or_add_trPr()
header_marker = OxmlElement("w:tblHeader")
header_marker.set(qn("w:val"), "true")
header_properties.append(header_marker)
headers = ["ສະຖານະ", "ຄວາມໝາຍ/ຂັ້ນຕອນຕໍ່ໄປ"]
for idx, text in enumerate(headers):
    cell = table.rows[0].cells[idx]
    cell.width = table.columns[idx].width
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), LIGHT)
    cell._tc.get_or_add_tcPr().append(shade)
    r = cell.paragraphs[0].add_run(text)
    set_font(r, 10, NAVY, bold=True)
rows = [
    ("ລໍຖ້າກວດ", "ໃບສັ່ງແລະຄິວຖືກສ້າງແລ້ວ; ລໍຖ້າພະນັກງານເອີ້ນຄິວ."),
    ("ເອີ້ນຄິວ", "ຄິວກຳລັງຖືກເອີ້ນ ແລະ ສະແດງໃນຈໍຄິວ."),
    ("ລໍຖ້າບັນທຶກຜົນ", "ກຳລັງດຳເນີນການກວດ; ຕ້ອງເຂົ້າໜ້າຜົນກວດ."),
    ("ຄ້າງຊຳລະ", "ມີຜົນກວດແລ້ວ ແລະ ພ້ອມຮັບເງິນ."),
    ("ຈ່າຍແລ້ວ", "ບັນທຶກ payment ສຳເລັດ ແລະ ສາມາດພິມໃບຮັບເງິນ."),
    ("ຍົກເລີກແລ້ວ", "ໃບສັ່ງບໍ່ດຳເນີນຕໍ່ ແຕ່ປະຫວັດຍັງຖືກເກັບ."),
]
for status, meaning in rows:
    cells = table.add_row().cells
    cells[0].width = Inches(1.7)
    cells[1].width = Inches(5.0)
    for cell, text in zip(cells, [status, meaning]):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(text), 9.6, None, bold=(cell is cells[0]))

doc.add_heading("ຂໍ້ຄວນລະວັງ", level=2)
add_bullets([
    "ກວດ HN, ຊື່ຄົນເຈັບ ແລະ ປະເພດກວດກ່ອນບັນທຶກຜົນ ຫຼື ຮັບເງິນ.",
    "ບໍ່ຄວນໃຊ້ປຸ່ມ Back/Refresh ຂອງ browser ໃນຂະນະທີ່ກຳລັງສົ່ງຟອມ.",
    "ການລົບ, Void ແລະ Refund ຕ້ອງກວດລາຍລະອຽດໃນ modal ແລະ ລະບຸເຫດຜົນໃຫ້ຊັດເຈນ.",
    "ຮູບຜົນກວດຮັບສະເພາະ JPG/PNG ບໍ່ເກີນ 5 MB.",
    "ເມື່ອສິ້ນສຸດການໃຊ້ງານ ໃຫ້ກົດ “ອອກຈາກລະບົບ” ເພື່ອລຶບ session cookie.",
])
add_note("ຂໍ້ມູນໃນຮູບໜ້າຈໍເປັນຂໍ້ມູນຕົວຢ່າງຈາກຖານຂໍ້ມູນທົດສອບ. ເມື່ອນຳໃຊ້ຈິງ ຕ້ອງຮັກສາຄວາມລັບຂອງຂໍ້ມູນຄົນເຈັບ.")

# Keep image captions with images and reduce widows.
for paragraph in doc.paragraphs:
    paragraph.paragraph_format.widow_control = True

OUT_DIR.mkdir(parents=True, exist_ok=True)
doc.save(OUT_FILE)
print(OUT_FILE.name.encode("unicode_escape").decode("ascii"))
